from fastapi import FastAPI, Request
from pydantic import BaseModel
from typing import Optional
import openai
import os
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from docx import Document
import tempfile

openai.api_key = os.getenv("OPENAI_API_KEY")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["https://lcardona21.github.io"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

class Proyecto(BaseModel):
    titulo: str
    grupo: str
    eje: str
    objetivo_general: str
    objetivo_esp_1: Optional[str] = None
    objetivo_esp_2: Optional[str] = None
    objetivo_esp_3: Optional[str] = None
    justificacion: Optional[str] = None
    descripcion: Optional[str] = None
    duracion: Optional[str] = None
    fechas: Optional[str] = None
    presupuesto_transporte: Optional[str] = None
    presupuesto_materiales: Optional[str] = None
    presupuesto_refrigerios: Optional[str] = None
    presupuesto_total: Optional[str] = None
    poblacion: Optional[str] = None
    finalidad: Optional[str] = None
    indicador_1: Optional[str] = None
    indicador_2: Optional[str] = None
    indicador_3: Optional[str] = None

@app.post("/generar_proyecto")
async def generar(data: Proyecto):
    prompt = f"""
Redacta un proyecto scout con los siguientes datos:

Título: {data.titulo}
Grupo: {data.grupo}
Eje: {data.eje}
Objetivo General: {data.objetivo_general}
Objetivos Específicos:
- {data.objetivo_esp_1 or ''}
- {data.objetivo_esp_2 or ''}
- {data.objetivo_esp_3 or ''}
Justificación: {data.justificacion or ''}
Descripción: {data.descripcion or ''}
Duración: {data.duracion or ''} ({data.fechas or ''})
Presupuesto:
- Transporte: {data.presupuesto_transporte or ''}
- Materiales: {data.presupuesto_materiales or ''}
- Refrigerios: {data.presupuesto_refrigerios or ''}
- Total: {data.presupuesto_total or ''}
Población objetivo: {data.poblacion or ''}
Finalidad: {data.finalidad or ''}
Indicadores de éxito:
- {data.indicador_1 or ''}
- {data.indicador_2 or ''}
- {data.indicador_3 or ''}

Usa lenguaje scout, claro, estructurado y motivador.
"""
    try:
        response = openai.ChatCompletion.create(
            model="gpt-4",
            messages=[{"role": "user", "content": prompt}]
        )
        proyecto_text = response.choices[0].message.content
        return JSONResponse(content={"proyecto": proyecto_text})
    except Exception as e:
        return JSONResponse(status_code=500, content={"error": str(e)})

@app.post("/descargar_docx")
async def descargar_docx(request: Request):
    payload = await request.json()
    texto = payload.get("proyecto", "")
    doc = Document()

    for bloque in texto.strip().split("\n\n"):
        if ":" in bloque:
            titulo, contenido = bloque.split(":", 1)
            table = doc.add_table(rows=2, cols=1)
            table.style = 'Table Grid'
            table.cell(0, 0).text = titulo.strip()
            table.cell(1, 0).text = contenido.strip()
            doc.add_paragraph()
        else:
            doc.add_paragraph(bloque.strip())

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    tmp.close()
    return FileResponse(
        path=tmp.name,
        filename="Proyecto_Rover.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
